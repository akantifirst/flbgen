from flask import Flask, render_template, request, send_file, redirect, url_for
import sqlite3
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your_secret_key_here'  # Замените на ваш секретный ключ

DATABASE = 'database.db'


def get_text_blocks(labels):
    if not labels:
        return []
    conn = sqlite3.connect(DATABASE)
    cursor = conn.cursor()
    placeholder = ','.join(['?'] * len(labels))
    query = f"SELECT content FROM text_blocks WHERE label IN ({placeholder})"
    cursor.execute(query, labels)
    results = cursor.fetchall()
    conn.close()
    return [row[0] for row in results]


def generate_docx(rooms):
    document = Document('template.docx')
    room_number = 1  # Счётчик для комнат
    for room in rooms:
        # Добавляем название помещения с нумерацией
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"1.{room_number} {room['name']}")
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        r = run._element.rPr.rFonts
        r.set(qn('w:eastAsia'), 'Arial')

        # Добавляем пустую строку для разделения
        document.add_paragraph()

        # Счётчик для компонентов
        feature_number = 1
        if room['features']:
            for feature in room['features']:
                # Получаем контент для feature
                content = get_text_blocks([feature])[0]

                # Добавляем нумерацию и название feature на отдельной строке
                paragraph = document.add_paragraph()
                run = paragraph.add_run(f"1.{room_number}.{feature_number} {feature}")
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                r = run._element.rPr.rFonts
                r.set(qn('w:eastAsia'), 'Arial')

                # Добавляем пустую строку для разделения
                document.add_paragraph()

                # Добавляем контент feature на отдельной строке под названием
                paragraph = document.add_paragraph(content)
                run = paragraph.runs[0]
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                r = run._element.rPr.rFonts
                r.set(qn('w:eastAsia'), 'Arial')

                # Добавляем пустую строку для разделения
                document.add_paragraph()

                feature_number += 1
        else:
            # Если у комнаты нет компонентов, выводим сообщение
            paragraph = document.add_paragraph(f"1.{room_number} Keine Komponenten ausgewählt.")
            run = paragraph.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            r = run._element.rPr.rFonts
            r.set(qn('w:eastAsia'), 'Arial')

        # Добавляем пустую строку для разделения комнат
        document.add_paragraph()

        room_number += 1

    # Сохраняем документ в буфер
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


@app.route('/', methods=['GET', 'POST'])
def index():
    conn = sqlite3.connect(DATABASE)
    cursor = conn.cursor()
    cursor.execute("SELECT label FROM text_blocks")
    labels = [row[0] for row in cursor.fetchall()]
    conn.close()

    if request.method == 'POST':
        room_types = request.form.getlist('room_type')
        # Получаем список всех возможных имен полей для функций
        feature_fields = [key for key in request.form.keys() if key.startswith('features_')]
        feature_fields_sorted = sorted(feature_fields, key=lambda x: int(x.split('_')[1]))

        rooms = []
        for idx, room_type in enumerate(room_types):
            features = request.form.getlist(f'features_{idx}')
            rooms.append({
                'name': room_type.strip(),
                'features': features
            })

        # Проверка ввода
        if not all(room['name'] for room in rooms):
            return render_template('index.html', labels=labels, error="Bitte alle Namen der Raumtypen eintragen.")

        # Генерация .docx контента
        buffer = generate_docx(rooms)

        return send_file(
            buffer,
            as_attachment=True,
            download_name='FLB.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    return render_template('index.html', labels=labels)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
